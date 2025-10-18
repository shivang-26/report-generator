from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document(project_data: dict) -> Document:
    """Create a new Word document with project data"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(project_data['title'], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add authors
    authors = doc.add_paragraph()
    authors.add_run("By: " + ", ".join(project_data['authors']))
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    from datetime import datetime
    date = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add sections
    sections = ['abstract', 'introduction', 'methodology', 'results', 'conclusion']
    for section in sections:
        if project_data.get(section):
            doc.add_heading(section.capitalize(), level=2)
            doc.add_paragraph(project_data[section])
    
    # Add references if available
    if project_data.get('references'):
        doc.add_heading('References', level=2)
        for ref in project_data['references']:
            doc.add_paragraph(ref['citation'], style='List Bullet')
    
    return doc
